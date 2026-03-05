import streamlit as st
import json
import os
import pandas as pd
from io import BytesIO
from docx import Document
from datetime import datetime

# Import backend logic
try:
    from backend import run_cdd_agent
except ImportError:
    st.error("Please ensure 'backend.py' is in the same folder.")
    st.stop()

# --- CONFIGURATION & STYLING ---
st.set_page_config(page_title="Adverse Media Screener", layout="wide", page_icon= "🚀")

# Custom CSS for EY-style Professional Look
st.markdown("""
    <style>
    /* Main Background and Fonts */
    .stApp {
        background-color: #fcfcfc;
        font-family: "Segoe UI", "Helvetica Neue", Arial, sans-serif;
    }
    
    /* Header Styling */
    h1 {
        color: #2E2E38;
        font-weight: 700;
        padding-bottom: 10px;
        border-bottom: 3px solid #FFE600; /* EY Yellow accent */
    }
    
    /* Section Headers (Scopes) */
    .scope-header {
        background-color: #2E2E38;
        color: white;
        padding: 8px 15px;
        border-radius: 4px;
        margin-top: 20px;
        margin-bottom: 15px;
        font-weight: 600;
        font-size: 1.1em;
        border-left: 5px solid #FFE600;
    }
    
    /* Button Styling */
    div.stButton > button:first-child {
        background-color: #2E2E38;
        color: white;
        border-radius: 4px;
        border: none;
        padding: 10px 24px;
        font-weight: 600;
    }
    div.stButton > button:first-child:hover {
        background-color: #4a4a55;
        border-color: #FFE600;
        color: white;
    }
    
    /* Custom Badge Class */
    .risk-badge {
        background-color: #ffe6e6;
        color: #b30000;
        padding: 4px 8px;
        border-radius: 4px;
        font-size: 0.85em;
        font-weight: 600;
        border: 1px solid #ffcccc;
    }
    
    /* Steps/Instruction Box */
    .instruction-box {
        background-color: #eef4ff;
        border-left: 5px solid #0055aa;
        padding: 15px;
        margin-bottom: 20px;
        border-radius: 4px;
    }
    </style>
""", unsafe_allow_html=True)

# --- HELPER FUNCTIONS ---

def generate_word_report(entity_name, selected_stories, summary_text):
    """Generates a Word document object for download."""
    doc = Document()
    heading = doc.add_heading(f'Adverse Media Report: {entity_name}', 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Confidentiality: Internal Use Only")
    
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(summary_text)
    
    doc.add_heading('Detailed Findings', level=1)
    
    if not selected_stories:
        doc.add_paragraph("No specific red flag items were selected for this report.")
    
    for story in selected_stories:
        # Sanitize headline
        headline = story.get('headline', 'Untitled Story').replace("_", " ")
        p = doc.add_heading(headline, level=2)
        
        # Red Flags
        flags = story.get('matched_flags', [])
        if flags:
            p = doc.add_paragraph()
            run = p.add_run("Identified Risks:")
            run.bold = True
            for flag in flags:
                doc.add_paragraph(f"• [{flag.get('code')}] {flag.get('name')}", style='List Bullet')
        
        # Summary
        doc.add_paragraph("Summary:", style='Heading 3')
        # Sanitize summary for Word
        clean_summary = story.get('summary', '').replace("_", " ")
        doc.add_paragraph(clean_summary)
        
        # Sources
        doc.add_paragraph("Sources:", style='Heading 3')
        for source in story.get('sources', []):
            doc.add_paragraph(f"{source.get('source_name')} - {source.get('published_date')}", style='List Bullet')
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def save_result_to_storage(entity_name, metadata, results):
    STORAGE_FILE = "screening_history.json"
    record = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "entity": entity_name,
        "metadata": metadata,
        "results": results
    }
    if os.path.exists(STORAGE_FILE):
        with open(STORAGE_FILE, "r") as f:
            try: history = json.load(f)
            except: history = []
    else:
        history = []
    history.append(record)
    with open(STORAGE_FILE, "w") as f:
        json.dump(history, f, indent=4)

def load_history():
    STORAGE_FILE = "screening_history.json"
    if not os.path.exists(STORAGE_FILE):
        return []
    with open(STORAGE_FILE, "r") as f:
        try:
            return json.load(f)
        except:
            return []

def clean_metadata_input(df):
    """Cleans the data editor output to prevent None/Empty rows."""
    if df.empty:
        return {}
    clean_df = df.dropna(subset=['Key'])
    clean_df = clean_df[clean_df['Key'].astype(str).str.strip() != '']
    clean_df['Value'] = clean_df['Value'].fillna("")
    return dict(zip(clean_df["Key"], clean_df["Value"]))


# --- APP LOGIC ---

# Sidebar Navigation
with st.sidebar:
    st.markdown("### **ZispireTech** | CDD Screener")
    page = st.radio("Navigate", ["New Screening", "Search History"])
    st.markdown("---")

    if page == "New Screening":
        st.subheader("1. Entity Details")
        entity_name_input = st.text_input("Target Entity Name", value="")
        
        st.subheader("2. Context & Metadata")
        st.caption("ℹ️ Add rows as needed. Rows with empty 'Keys' are ignored.")
        
        default_data = pd.DataFrame([
            {"Key": "Country", "Value": ""},
            {"Key": "Sector", "Value": ""},
        ])
        
        edited_df = st.data_editor(
            default_data, 
            num_rows="dynamic", 
            hide_index=True, 
            use_container_width=True,
            column_config={
                "Key": st.column_config.TextColumn("Parameter Name", required=True),
                "Value": st.column_config.TextColumn("Value")
            }
        )
        
        metadata_input = clean_metadata_input(edited_df)
        st.markdown("---")
        run_btn = st.button("Run Screening Analysis", type="primary")

# --- PAGE 1: NEW SCREENING ---
if page == "New Screening":
    col1, col2 = st.columns([0.8, 0.2])
    with col1:
        st.title("Adverse Media Screening")
        st.markdown("**Automated Red Flag Analysis & Risk Summarization**")

    if "screening_results" not in st.session_state:
        st.session_state.screening_results = None

    if run_btn:
        if not entity_name_input:
            st.warning("⚠️ Please enter an entity name to begin.")
        else:
            with st.spinner("Consulting sources and analyzing risks..."):
                results = run_cdd_agent(entity_name=entity_name_input, metadata=metadata_input)
                st.session_state.screening_results = results
                save_result_to_storage(entity_name_input, metadata_input, results)
                st.rerun()

    # Display Results
    if st.session_state.screening_results:
        results = st.session_state.screening_results
        
        st.markdown("---")
        st.subheader(f"Results for: {entity_name_input}")
        
        # INSTRUCTIONS (New Feature)
        st.markdown("""
        <div class="instruction-box">
            <strong>📋 How to generate a report:</strong><br>
            1. Review the red flags below.<br>
            2. <strong>Tick the checkbox</strong> next to any story/allegation you want to include.<br>
            3. Scroll to the bottom and click "Generate Summary Report".
        </div>
        """, unsafe_allow_html=True)

        with st.form("report_generation_form"):
            selected_stories = []
            
            for i, scope_data in enumerate(results):
                scope_name = scope_data.get("scope", "Unknown Scope")
                stories = scope_data.get("stories", [])
                
                st.markdown(f"<div class='scope-header'>{scope_name}</div>", unsafe_allow_html=True)
                
                if not stories:
                    st.caption("No adverse media found in this category.")
                    continue

                for idx, story in enumerate(stories):
                    with st.container():
                        c1, c2 = st.columns([0.05, 0.95])
                        
                        # Checkbox
                        with c1:
                            st.write("") 
                            is_selected = st.checkbox("", key=f"sel_{i}_{idx}")
                            if is_selected:
                                selected_stories.append(story)
                        
                        # Content
                        with c2:
                            headline_clean = story.get('headline', 'Untitled').replace("_", " ")
                            st.markdown(f"#### {headline_clean}")
                            
                            flags = story.get("matched_flags", [])
                            if any(f.get('code') for f in flags):
                                flag_html = ""
                                for f in flags:
                                    if f.get('code'):
                                        flag_html += f"<span class='risk-badge'>🚩 {f['name']} ({f['code']})</span> "
                                st.markdown(flag_html, unsafe_allow_html=True)
                                
                                for f in flags:
                                    if f.get('rationale'):
                                        clean_rationale = f.get('rationale', '').replace("_", " ")
                                        st.caption(f"**Analysis:** {clean_rationale}")
                            
                            clean_summary = story.get("summary", "").replace("_", " ")
                            st.write(clean_summary)
                            
                            with st.expander("View Sources"):
                                for src in story.get("sources", []):
                                    st.markdown(f"- [{src.get('source_name')}]({src.get('url')}) - *{src.get('published_date')}*")
                        
                        st.divider()

            st.markdown("### 📝 Report Generation")
            col_submit, col_download = st.columns([1, 1])
            with col_submit:
                generate_click = st.form_submit_button("Generate Summary Report")

        if generate_click:
            if not selected_stories:
                st.error("Please select at least one item to generate a report.")
            else:
                risk_areas = list(set([s.get('matched_flags')[0]['code'] for s in selected_stories if s.get('matched_flags')]))
                summary_text = (
                    f"Based on the analysis of {len(selected_stories)} adverse media articles, "
                    f"risk exposure has been identified in areas related to {', '.join(risk_areas) if risk_areas else 'general reputational risk'}. "
                    "Immediate escalation to the compliance committee is recommended."
                )
                
                st.success("Report Generated Successfully")
                st.info(f"**Executive Summary:** {summary_text}")
                
                docx_file = generate_word_report(entity_name_input, selected_stories, summary_text)
                
                st.download_button(
                    label="📄 Download Report as .docx",
                    data=docx_file,
                    file_name=f"CDD_Report_{entity_name_input}_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# --- PAGE 2: HISTORY ---
elif page == "Search History":
    st.title("Search History Log")
    
    history_data = load_history()
    
    if not history_data:
        st.info("No search history found yet.")
    else:
        # Prepare data for cleaner display
        table_rows = []
        for entry in history_data:
            # Calculate total stories found
            total_stories = 0
            risk_scopes = set()
            if entry.get("results"):
                for r in entry["results"]:
                    count = len(r.get("stories", []))
                    total_stories += count
                    if count > 0:
                        risk_scopes.add(r.get("scope"))
            
            table_rows.append({
                "Timestamp": entry.get("timestamp"),
                "Entity Name": entry.get("entity"),
                "Metadata": str(entry.get("metadata")), # Flatten dict for display
                "Stories Found": total_stories,
            })
            
        df_history = pd.DataFrame(table_rows)
        
        # Sort by latest first
        df_history = df_history.sort_values(by="Timestamp", ascending=False)
        
        st.dataframe(
            df_history, 
            use_container_width=True, 
            hide_index=True,
            column_config={
                "Timestamp": st.column_config.TextColumn("Date & Time", width="medium"),
                "Entity Name": st.column_config.TextColumn("Entity", width="medium"),
                "Stories Found": st.column_config.NumberColumn("Hits", width="small"),
            }
        )